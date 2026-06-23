"""Build Project_Case_Study.docx and Interview_Talking_Points.docx"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from pathlib import Path
BASE = Path(__file__).resolve().parents[1]
FOREST_GREEN = RGBColor(0x1B, 0x43, 0x32)
CHARCOAL     = RGBColor(0x2C, 0x3E, 0x50)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)

def dh(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    c = FOREST_GREEN if level <= 2 else CHARCOAL
    for run in h.runs:
        run.font.color.rgb = c
        run.font.size = Pt(16 if level==1 else 13 if level==2 else 11)

def dp(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = CHARCOAL
    return p

# ── CASE STUDY ───────────────────────────────────────────────────────────────
print("Building Project_Case_Study.docx...")
doc = Document()
doc.add_heading("PROJECT CASE STUDY", 0).runs[0].font.color.rgb = FOREST_GREEN
doc.add_heading("Parastatal X Dynamic Budget Monitoring, Forecasting and Management Dashboard", 1).runs[0].font.color.rgb = FOREST_GREEN
dp(doc, "STAR Method Case Study | Portfolio Project | All data is entirely synthetic", italic=True)
dp(doc, "")

dh(doc, "SITUATION", 1)
dp(doc, """Many Kenyan State Corporations face a persistent challenge in financial management: 
budget data exists in silos across multiple systems — manual spreadsheets, legacy IFMIS 
reports, and paper-based registers — making it extremely difficult for management to obtain 
a consolidated, timely view of their financial position.

The consequences are significant. By the time reports reach the Chief Executive Officer or 
the Board, the data is weeks old. Expenditure variances go undetected until the financial 
year is nearly over. Approval process bottlenecks accumulate quietly until they block 
entire cost centres from spending. Audit findings repeat year after year because there is 
no systematic way to track their resolution.

I set out to design a financial analytics solution that addresses these challenges in the 
Kenyan public sector context — incorporating the specific regulatory frameworks, budget 
structures, and reporting requirements that define how Kenyan parastatals operate.""")

dh(doc, "TASK", 1)
dp(doc, """My assignment was to design, build, and document a complete financial analytics 
project for a fictional Kenyan State Corporation — demonstrating the full analytics 
lifecycle from business understanding through to portfolio presentation.

The deliverables included:
• A realistic synthetic dataset representing three financial years of operations
• A professionally formatted Excel workbook with 26 worksheets
• A star-schema data model suitable for Power BI
• Complete Power Query M transformation scripts
• A library of 70+ DAX measures
• A financial forecasting model using four methods
• A 12-page Power BI dashboard design
• Senior Finance Manager management commentary
• A management action plan and financial risk register
• LinkedIn and GitHub portfolio materials
• A 15-slide executive presentation

The project had to be technically rigorous, financially credible, contextually accurate for 
the Kenyan public sector, and suitable for presentation to senior management, recruiters, 
and professional networks.""")

dh(doc, "ACTION", 1)

dh(doc, "1. Synthetic Data Generation", 2)
dp(doc, """I wrote a comprehensive Python script (generate_synthetic_data.py) using Pandas, 
NumPy and the random module with a fixed seed (42) for reproducibility. The script generates 
18,896+ rows across 13 datasets, including:

• Approved annual budgets across 32 cost centres, 13 directorates and 8 regions
• Monthly budget allocation profiles with realistic seasonality (Jul-Jun weights)
• Expenditure transactions with seasonal spend patterns — higher in Q3/Q4 and June
• Commitment registers with multiple lifecycle statuses
• Pending bills with aging profiles and dispute tracking
• Revenue collections with banking delay tracking
• Performance indicators with 15 KPIs across 3 financial years
• Audit exceptions with repeat-finding tracking
• Intentional data quality issues at a 2-5% rate for cleaning demonstration

All amounts are in Kenya Shillings. Budget scales are realistic for a mid-sized parastatal 
(KSh 1.70 billion in FY2025/2026). Financial years run from July 1 to June 30.""")

dh(doc, "2. Data Validation and Reconciliation", 2)
dp(doc, """I wrote automated validation scripts running 36 checks across all datasets, 
including primary key uniqueness, financial year coverage, budget-to-allocation reconciliation, 
negative amount detection, and amount consistency checks. All 36 checks pass.

A reconciliation script performs 18 balance checks confirming that monthly allocations sum 
to annual budgets (zero difference), pending bill invoices equal paid plus outstanding amounts, 
and revenue records self-reconcile. Variance-explained records are documented.""")

dh(doc, "3. Data Cleaning and Governance", 2)
dp(doc, """I introduced controlled data quality issues at approximately 2.6% — consistent 
with the 2-5% target — including missing codes, duplicate transaction references, inconsistent 
region names, negative amounts, and invalid dates. These are documented in a 250-row 
Data_Quality_Log.csv with original values, corrected values, correction rules, severity 
ratings, and resolution status.

Power Query M scripts apply systematic cleaning: text trimming, case standardisation, 
code validation, duplicate flagging, financial year month calculation, and DateKey creation 
for relationship building.""")

dh(doc, "4. Star Schema Data Model", 2)
dp(doc, """I designed a star schema with 12 fact tables and 14 dimension tables. The single 
date spine (Dim_Date) connects to all fact tables via DateKey (YYYYMMDD integer). 
Financial year months are numbered July=1 through June=12 to support correct FY-relative 
calculations.

Design principles applied:
• No many-to-many relationships
• Bidirectional filtering disabled by default
• Approved budget separated from revised/working budget
• Commitments and actual expenditure tracked separately to prevent double-counting
• RAG thresholds stored in a configurable separate table""")

dh(doc, "5. DAX Measure Development", 2)
dp(doc, """I developed 70+ DAX measures organised into display folders:
• Budget measures: Approved, Revised, YTD, Balance, Utilisation %, Absorption %
• Expenditure: Total, Monthly, YTD, Average, Run Rate, YoY Growth, QoQ Movement
• Commitments: Open, Outstanding, Total Exposure, Available Balance after Commits
• Requests: Count, Pending, Approval Rate, Processing Rate, Aging (30/60 days)
• Revenue: Budgeted, Actual, Achievement %, Unreconciled, Banking Delays
• Performance: Targets, Achievement %, Cost per Output, Budget Efficiency Ratio
• Risk: Negative Balance Count, Over-Budget CCs, Low Absorption CCs, Exception Count
• Forecasting: 4-method forecasts, Blended, Remaining Months, Required Monthly Spend
• RAG Status: Absorption, Forecast, Bills, Data Quality
• Control: Last Refresh Date, Reconciliation Difference, DQ Score

All measures use DIVIDE() for safe division and appropriate error handling.""")

dh(doc, "6. Financial Forecasting", 2)
dp(doc, """I implemented four year-end forecasting methods and a blended selection:

Method 1 – Straight-line run rate (YTD ÷ months elapsed × 12): KSh 1,612,762,691
Method 2 – 3-month weighted moving average: KSh 1,653,229,667
Method 3 – Seasonal adjustment (historical July-June pattern): KSh 1,722,723,783
Method 4 – Commitment-adjusted (YTD + open commitments): KSh 2,178,365,800

Blended forecast (30%/25%/25%/20%): KSh 1,763,490,330 (103.7% absorption)
Over-expenditure risk: KSh 63,490,330

Three scenarios are documented: Base Case, Optimistic, and Adverse, with recommended 
management responses for each.""")

dh(doc, "7. Dashboard Design", 2)
dp(doc, """I designed 12 Power BI dashboard pages with specifications, visual types, and 
layout instructions documented in the PowerBI_Implementation_Guide.docx. A custom theme 
(Theme_JSON.json) applies the forest green / muted gold / white / charcoal palette.

Key design decisions:
• Executive overview with waterfall chart (Budget → Expenditure → Balance)
• 4-quadrant scatter plot for performance and value-for-money analysis
• Drill-through from any summary page to cost-centre detail
• Configurable RAG conditional formatting via threshold table
• Management commentary text box on Page 1
• Last refresh date and data quality score prominently displayed""")

dh(doc, "RESULT", 1)
dp(doc, """The completed project demonstrates the following capabilities and potential benefits:

IMPROVED BUDGET VISIBILITY
The dashboard consolidates data from 13 datasets into a single view, allowing management 
to see the complete financial picture — approved budget, YTD spend, commitments, pending 
bills, revenue, and forecasts — in one place rather than across multiple disparate reports.

FASTER EXCEPTION DETECTION
The RAG traffic light system flags 23 cost centres as Red (below absorption expectations) 
and highlights 5 vote lines at over-expenditure risk, enabling management to intervene 
months before the financial year closes rather than discovering problems in the annual audit.

STRONGER EXPENDITURE CONTROL
The available-balance-after-commitments measure gives management a more conservative and 
accurate picture of available funds, incorporating KSh 700M in open commitments that would 
not appear in a simple budget-minus-expenditure calculation.

BETTER APPROVAL MONITORING
The AIE workflow page reveals a KSh 2.48B backlog of unprocessed approved requests — a 
systemic bottleneck that was invisible in conventional reporting. The aging analysis 
identifies how long each request has been outstanding, enabling targeted management action.

IMPROVED FORECASTING
The 4-method blended forecast model provides management with a range of year-end scenarios 
rather than a single optimistic projection, supporting more informed and conservative 
decision-making.

BETTER MANAGEMENT ACCOUNTABILITY
The management action register provides a live tracking tool for all identified financial 
risks, with responsible roles, target dates, progress percentages and follow-up comments.

STRONGER AUDIT READINESS
The audit exceptions dashboard tracks all findings by type, risk rating, status and repeat 
occurrence — creating a structured mechanism for management response that supports both 
internal review and external audit preparation.

PORTFOLIO VALUE
The project demonstrates to recruiters and interview panels proficiency across the full 
analytics stack: Python data engineering, data governance, star schema design, DAX 
development, Power Query, financial forecasting, management reporting, and executive 
communication.""")

dp(doc, "")
dp(doc, "Prepared by: Financial Analytics Portfolio Project", italic=True)
dp(doc, "Date: June 2026 | All data synthetic | Not representative of any real institution", italic=True)

doc.save(f"{BASE}/05_Portfolio/Project_Case_Study.docx")
print("  Project_Case_Study.docx saved.")

# ── INTERVIEW TALKING POINTS ─────────────────────────────────────────────────
print("Building Interview_Talking_Points.docx...")
doc2 = Document()
doc2.add_heading("PARASTATAL X – INTERVIEW TALKING POINTS", 0).runs[0].font.color.rgb = FOREST_GREEN
dp(doc2, "Prepared for: Financial Analytics Portfolio Project Interviews", italic=True, bold=True)
dp(doc2, "")

talking_points = [
    ("Why did you develop this project?",
     """I wanted to demonstrate how modern data analytics tools can be applied to the specific 
challenges of Kenyan public sector financial management. Most public sector finance 
demonstrations use generic data. I wanted to show genuine contextual knowledge — the July-June 
financial year, vote code structures, AIE workflows, supplementary budgets, AIA revenue, 
and donor fund tracking — alongside strong technical delivery."""),

    ("Why did you use synthetic data?",
     """Using synthetic data was both ethically correct and strategically sound. Government 
financial data is confidential. Using real data without authorisation would be a data 
governance breach. Synthetic data allowed me to generate a realistic, high-volume, 
internally consistent dataset that demonstrates all the scenarios I needed — over-expenditure 
risk, pending bill backlogs, commitment exposure, audit exceptions — without any risk of 
misrepresenting a real institution. I used Python's random module with seed=42, so the 
data is fully reproducible."""),

    ("How did you incorporate the Kenyan public-sector context?",
     """I embedded several Kenya-specific features throughout. The financial year runs July 
1 to June 30 as per Kenyan law. I modelled Authority to Incur Expenditure (AIE) workflows, 
vote code structures matching the Kenya Public Finance Management Act classifications, 
supplementary budget mechanics, appropriations-in-aid revenue streams, donor fund codes, 
performance contracting KPIs, and the quarterly M&E reporting cycle. Budget amounts are in 
Kenya Shillings at a scale realistic for a mid-size state corporation."""),

    ("How were the datasets generated?",
     """I used Python with Pandas and NumPy. The script is approximately 450 lines and 
produces 13 CSV files plus 14 dimension tables in a single run. Key design choices included 
seasonal expenditure weights (higher in Q3/Q4 and particularly June), realistic procurement 
and commitment lifecycles, approval processing bottlenecks, and an intentional 2.6% data 
quality issue rate for cleaning demonstration. The random seed ensures the data is 
reproduced identically on every run."""),

    ("How did you manage data quality?",
     """I introduced 18 types of intentional data quality issues at approximately 2.6% 
prevalence — including missing codes, duplicate references, negative amounts, invalid dates, 
inconsistent names and capitalisation. These are all documented in a 250-row Data Quality Log 
with original values, corrected values, correction rules and severity ratings.

In Power Query, I applied systematic cleaning: Text.Trim() for whitespace, Text.Upper() 
for code standardisation, date validation, duplicate detection, and DataQualityFlag columns 
to segregate clean from problematic records. The model distinguishes Clean, Warning, and 
Error records throughout."""),

    ("Why did you choose a star schema?",
     """Star schemas are optimised for Power BI DAX calculations. They minimise many-to-many 
relationships, ensure clean filter propagation from dimension to fact, and significantly 
improve query performance compared to flat tables or snowflake designs. 

For this project, the single Dim_Date table connected via DateKey (YYYYMMDD integer) provides 
the time intelligence spine, while shared dimensions like Dim_Cost_Centre, Dim_Vote and 
Dim_Region serve multiple fact tables simultaneously without duplication."""),

    ("How were DAX measures designed?",
     """I organised 70+ measures into display folders by category — budget, expenditure, 
commitment, requests, revenue, performance, risk, forecasting, RAG status, and control. 
Every measure uses DIVIDE() instead of the division operator to handle zero denominators 
safely. Time intelligence measures use DATESYTD() with 30/6 as the financial year end 
date. YoY measures use SAMEPERIODLASTYEAR(). Forecast measures chain from YTD actuals 
using the remaining months calculation."""),

    ("How were forecasts calculated?",
     """Four methods were implemented: straight-line run rate, three-month moving average, 
seasonal historical adjustment, and commitment-adjusted forecast. Each has different 
strengths — the commitment method captures crystallisation risk while the seasonal method 
reflects year-end spending patterns.

I blended them at 30/25/25/20 weights to balance between conservative and aggressive 
estimates. The blended result of KSh 1.763B against a KSh 1.700B budget creates an 
interesting management scenario — slight over-expenditure risk requiring action."""),

    ("How were management risks identified?",
     """Every finding in the management commentary is derived directly from the synthetic 
data — not generic assertions. For example: '23 cost centres rated Red' comes from the 
actual RAG calculation on 32 cost centres; 'KSh 63.5M over-expenditure risk' comes from 
the blended forecast minus approved budget; 'KSh 2.48B unprocessed requests' comes from 
the requests table filtered by status and outstanding amounts. The discipline of grounding 
every observation in data is something I applied throughout."""),

    ("How does the dashboard support decision-making?",
     """The dashboard is structured around the management decision cycle. The Executive 
Overview gives the CEO a 30-second read. The drill-through pages allow the Director of 
Finance to investigate specific cost centres or vote lines. The workflow page shows the 
budget team their processing backlog. The risk register gives the Head of Internal Audit 
a live exception tracking tool. Each page is designed for a specific audience and decision 
type — not just data display."""),

    ("What challenges did you encounter?",
     """Several interesting ones. Getting the expenditure amounts to scale correctly relative 
to the budget required careful calibration — my first attempt generated KSh 18B in 
expenditure against a KSh 1.7B budget. The Excel workbook initially timed out when writing 
12,100 transaction rows with openpyxl, so I switched to XlsxWriter which is dramatically 
faster. Calibrating the commitment amounts to create a realistic but not absurd exposure 
required multiple iterations."""),

    ("What improvements could be made?",
     """Several natural extensions: connecting to a real IFMIS data extract (with appropriate 
authorisation) to make this operational; adding Power Automate flows to trigger alerts when 
RAG thresholds are breached; building mobile-optimised report pages for Regional Managers; 
adding NLP-generated management commentary using Azure OpenAI; and connecting to the Kenya 
National Treasury API if one becomes available for benchmark data."""),

    ("How could this be deployed in a real institution?",
     """The architecture maps cleanly to a real deployment. The CSV files would be replaced 
by scheduled IFMIS exports or a direct database connection. Power BI Service would host the 
dashboard with role-based access — the CEO sees the executive overview, cost centre managers 
see only their own data via Row Level Security, the Finance Director sees everything. The 
monthly refresh would be automated via Power Automate or a scheduled gateway connection. 
The RAG threshold table allows non-technical managers to adjust tolerance levels without 
touching the code."""),
]

for q, a in talking_points:
    dh(doc2, q, 2)
    dp(doc2, a)
    dp(doc2, "")

doc2.save(f"{BASE}/05_Portfolio/Interview_Talking_Points.docx")
print("  Interview_Talking_Points.docx saved.")
print("Portfolio docs built.")
