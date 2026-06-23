"""
Parastatal X - Document Builder
Creates all DOCX files for the project.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

from pathlib import Path
BASE = Path(__file__).resolve().parents[1]

FOREST_GREEN = RGBColor(0x1B, 0x43, 0x32)
MUTED_GOLD   = RGBColor(0xB7, 0x95, 0x0B)
CHARCOAL     = RGBColor(0x2C, 0x3E, 0x50)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)

def set_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color or FOREST_GREEN
        if level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(13)
        else:
            run.font.size = Pt(11)

def add_para(doc, text, bold=False, italic=False, size=10, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color or CHARCOAL
    return p

def add_table(doc, headers, rows, header_color=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = WHITE
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "1B4332")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shading)
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx+1].cells
        bg = "F2F2F2" if r_idx%2==0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            row_cells[c_idx].text = str(val)
            for run in row_cells[c_idx].paragraphs[0].runs:
                run.font.size = Pt(9)
            shading = OxmlElement("w:shd")
            shading.set(qn("w:fill"), bg)
            shading.set(qn("w:color"), "auto")
            shading.set(qn("w:val"), "clear")
            row_cells[c_idx]._tc.get_or_add_tcPr().append(shading)
    return table

# ═══════════════════════════════════════════════════════════════════════════
# DATA MODEL SPECIFICATION
# ═══════════════════════════════════════════════════════════════════════════
print("Building Data_Model_Specification.docx...")
doc = Document()
doc.add_heading("Parastatal X – Data Model Specification", 0)
doc.paragraphs[0].runs[0].font.color.rgb = FOREST_GREEN
doc.paragraphs[0].runs[0].font.size = Pt(20)

add_para(doc, "Dynamic Budget Monitoring, Forecasting and Management Dashboard", bold=True, size=12)
add_para(doc, "Version: 1.0  |  Date: June 2026  |  All data is synthetic", italic=True)
add_para(doc, "")

set_heading(doc, "1. Star Schema Architecture", 1)
add_para(doc, """
The data model uses a star schema design with a central set of fact tables 
connected to shared dimension tables. This design optimises Power BI DAX 
performance, ensures clean filter propagation, and avoids ambiguous 
many-to-many relationships.
""")

set_heading(doc, "1.1 Fact Tables", 2)
fact_table_rows = [
    ("Fact_Approved_Budget",     "96",   "Static approved budget by cost centre, vote and FY"),
    ("Fact_Monthly_Budget",      "1,152","Monthly budget allocation profiles (Jul–Jun)"),
    ("Fact_Supplementary_Budget","209",  "Supplementary budget adjustments"),
    ("Fact_Reallocation",        "224",  "Budget reallocation requests"),
    ("Fact_Expenditure",         "12,100","All expenditure transactions across 3 FYs"),
    ("Fact_Commitments",         "1,350","Outstanding commitment register"),
    ("Fact_Budget_Requests",     "1,000","AIE and budget review workflow"),
    ("Fact_Procurement",         "390",  "Procurement plan and contract status"),
    ("Fact_Pending_Bills",       "520",  "Outstanding supplier invoices"),
    ("Fact_Revenue",             "1,140","Revenue and Appropriations-in-Aid"),
    ("Fact_Performance",         "165",  "KPI tracking and M&E indicators"),
    ("Fact_Control_Exceptions",  "300",  "Internal audit and control findings"),
    ("Fact_Data_Quality",        "250",  "Data quality log"),
]
add_table(doc, ["Fact Table","Row Count","Description"], fact_table_rows)
add_para(doc, "")

set_heading(doc, "1.2 Dimension Tables", 2)
dim_rows = [
    ("Dim_Date",        "1,096 dates","Full date calendar 1 Jul 2023 – 30 Jun 2026","DateKey (INT)"),
    ("Dim_Financial_Year","3 FYs",   "FY2023/24, FY2024/25, FY2025/26","FinancialYear (TEXT)"),
    ("Dim_Programme",   "15",        "Programme hierarchy","ProgrammeCode"),
    ("Dim_Directorate", "13",        "13 directorates","DirectorateCode"),
    ("Dim_Cost_Centre", "32",        "HQ, regional, project and technical CCs","CostCentreCode"),
    ("Dim_Region",      "8",         "8 operational regions","RegionCode"),
    ("Dim_Activity",    "40",        "Activity catalogue","ActivityCode"),
    ("Dim_Sub_Activity","~150",      "Sub-activity catalogue","SubActivityCode"),
    ("Dim_Vote",        "33",        "Vote code (economic classification)","VoteCode"),
    ("Dim_Fund_Source", "8",         "GOK, AIA, donor sources","FundSourceCode"),
    ("Dim_Donor",       "8",         "Donor organisations","DonorCode"),
    ("Dim_Risk_Rating", "4",         "Critical, High, Medium, Low","RiskRating"),
    ("Dim_Officer_Role","29",        "Management and officer roles","OfficerRole"),
]
add_table(doc, ["Dimension","Rows","Description","Primary Key"], dim_rows)
add_para(doc, "")

set_heading(doc, "2. Key Relationships", 1)
rel_rows = [
    ("Fact_Expenditure","CostCentreCode","Dim_Cost_Centre","CostCentreCode","Many-to-One","Single"),
    ("Fact_Expenditure","VoteCode","Dim_Vote","VoteCode","Many-to-One","Single"),
    ("Fact_Expenditure","RegionCode","Dim_Region","RegionCode","Many-to-One","Single"),
    ("Fact_Expenditure","FundSourceCode","Dim_Fund_Source","FundSourceCode","Many-to-One","Single"),
    ("Fact_Expenditure","DateKey","Dim_Date","DateKey","Many-to-One","Single"),
    ("Fact_Expenditure","FinancialYear","Dim_Financial_Year","FinancialYear","Many-to-One","Single"),
    ("Fact_Approved_Budget","CostCentreCode","Dim_Cost_Centre","CostCentreCode","Many-to-One","Single"),
    ("Fact_Approved_Budget","VoteCode","Dim_Vote","VoteCode","Many-to-One","Single"),
    ("Fact_Monthly_Budget","CostCentreCode","Dim_Cost_Centre","CostCentreCode","Many-to-One","Single"),
    ("Fact_Commitments","CostCentreCode","Dim_Cost_Centre","CostCentreCode","Many-to-One","Single"),
    ("Fact_Revenue","RegionCode","Dim_Region","RegionCode","Many-to-One","Single"),
    ("Fact_Performance","DirectorateCode","Dim_Directorate","DirectorateCode","Many-to-One","Single"),
]
add_table(doc, ["From Table","From Column","To Table","To Column","Cardinality","Filter Direction"], rel_rows)
add_para(doc, "")

set_heading(doc, "3. Data Dictionary (Key Fields)", 1)
dd_rows = [
    ("ApprovedAnnualAmount","Number","KSh","Official approved budget for the financial year"),
    ("NetAmount","Number","KSh","Expenditure net of VAT (actual cash amount)"),
    ("OutstandingCommitment","Number","KSh","Commitment amount yet to be paid"),
    ("OutstandingAmount","Number","KSh","Invoice amount not yet paid (pending bills)"),
    ("AgingDays","Number","Days","Days elapsed since date received/commitment date"),
    ("FYMonthNumber","Number","1-12","Financial year month number (July=1, June=12)"),
    ("DataQualityFlag","Text","Clean/Warning/Error","Data quality status from Power Query"),
    ("RiskRating","Text","Critical/High/Medium/Low","Risk severity classification"),
    ("BudgetVersion","Text","Original/Revised","Distinguishes approved vs supplementary-adjusted"),
    ("CommitmentStatus","Text","Open/Partially Invoiced etc.","Lifecycle stage of commitment"),
    ("PaymentStatus","Text","Paid/Pending Payment/Cancelled","Payment lifecycle stage"),
    ("PerformanceStatus","Text","On Track/At Risk/Off Track","KPI performance classification"),
]
add_table(doc, ["Field","Type","Unit/Values","Description"], dd_rows)
add_para(doc, "")

set_heading(doc, "4. Model Design Rules", 1)
rules = [
    "Never use Fact_Approved_Budget as a cross-filter source into expenditure facts.",
    "The Dim_Date table is the single date spine — all date relationships go through DateKey.",
    "Financial Year is a text field (e.g. FY2025/2026) to avoid calendar-year confusion.",
    "FYMonthNumber (1=July, 12=June) is used for all time-intelligence month comparisons.",
    "The approved budget table is protected — never silently replace with revised figures.",
    "Commitments and expenditure are tracked separately; avoid double-counting in Total Exposure.",
    "All currency amounts are in full Kenya Shillings. KSh000 and KShM columns are derived.",
    "Data quality exceptions are stored in a separate fact table (Fact_Data_Quality).",
    "RAG thresholds are stored in a separate RAG_Thresholds query for configurability.",
    "Bidirectional filtering is disabled by default to prevent ambiguous results.",
]
for i, rule in enumerate(rules, 1):
    add_para(doc, f"{i}. {rule}")

set_heading(doc, "5. Implementation Checklist", 1)
checklist = [
    ("Import all 13 CSV files using Power Query","Done"),
    ("Apply M transformations for each query","Done"),
    ("Create all dimension tables","Done"),
    ("Build Dim_Date manually in Power Query","Done"),
    ("Set relationships as specified","Pending"),
    ("Import DAX measures from DAX_Measures.txt","Pending"),
    ("Apply Theme_JSON.json in Power BI View menu","Pending"),
    ("Design 12 dashboard pages per wireframes","Pending"),
    ("Configure RAG conditional formatting","Pending"),
    ("Enable drill-through on Cost Centre page","Pending"),
    ("Test all reconciliation measures","Pending"),
    ("Publish to Power BI Service (if applicable)","Pending"),
]
add_table(doc, ["Task","Status"], checklist)

filepath = f"{BASE}/03_PowerBI_Build_Pack/Data_Model_Specification.docx"
doc.save(filepath)
print(f"  Saved: {filepath}")

# ═══════════════════════════════════════════════════════════════════════════
# POWER BI IMPLEMENTATION GUIDE
# ═══════════════════════════════════════════════════════════════════════════
print("Building PowerBI_Implementation_Guide.docx...")
doc2 = Document()
doc2.add_heading("Parastatal X – Power BI Implementation Guide", 0)
doc2.paragraphs[0].runs[0].font.color.rgb = FOREST_GREEN
doc2.paragraphs[0].runs[0].font.size = Pt(18)

add_para(doc2, "Complete step-by-step guide for building the Power BI dashboard", bold=True, size=11)
add_para(doc2, "")

set_heading(doc2, "1. Prerequisites", 1)
prereqs = [
    "Power BI Desktop (latest version recommended — June 2026 or later)",
    "All 13 CSV source files in a consistent folder (e.g. C:\\ParastalX\\Data\\)",
    "DAX_Measures.txt, Power_Query_M_Code.txt, and Theme_JSON.json from the build pack",
    "Basic familiarity with Power Query and DAX",
]
for p in prereqs:
    add_para(doc2, f"• {p}")

set_heading(doc2, "2. Step-by-Step Build Instructions", 1)
steps = [
    ("Step 1: Create new PBIX file",
     "Open Power BI Desktop. File > New. Save as Parastatal_X_Dashboard.pbix."),
    ("Step 2: Apply theme",
     "View > Themes > Browse for themes > Select Theme_JSON.json. All colours will be applied automatically."),
    ("Step 3: Import data — Approved Budget",
     "Home > Get Data > Text/CSV > Select Approved_Annual_Budget.csv. Apply M code from Power_Query_M_Code.txt."),
    ("Step 4: Import all remaining CSV files",
     "Repeat for all 13 source files. Rename each query to match the Fact_/Dim_ naming in the data model."),
    ("Step 5: Build Dim_Date",
     "New Query > Blank Query > Paste Dim_Date M code from Power_Query_M_Code.txt. This generates all dates."),
    ("Step 6: Build RAG_Thresholds",
     "New Query > Blank Query > Paste RAG_Thresholds M code. This creates the configurable threshold table."),
    ("Step 7: Close and Apply",
     "Home > Close & Apply. All data loads into the model."),
    ("Step 8: Create relationships",
     "Model view > Drag fields to create relationships as specified in Data_Model_Specification.docx. Set cardinality and filter direction as documented."),
    ("Step 9: Import DAX measures",
     "Open DAX_Measures.txt. In Power BI, create a new table named [Budget Measures], [Expenditure Measures] etc. Paste each measure into the formula bar. Organise into display folders."),
    ("Step 10: Build Page 1 — Executive Overview",
     "Insert 10 KPI cards for main metrics. Add waterfall chart (Budget > Expenditure > Balance). Add monthly trend line chart. Add RAG table for top exceptions. Add last refresh card."),
    ("Step 11: Build remaining 11 pages",
     "Follow page specifications in Dashboard_Wireframes section. Use the measure names from DAX_Measures.txt for all visuals."),
    ("Step 12: Configure drill-through",
     "On Page 11 (Cost Centre drill-through): Right-click the CostCentreCode field > Add as drill-through field. All other pages will now link through to this detail page."),
    ("Step 13: Configure conditional formatting",
     "On all RAG indicators: Format > Conditional formatting > Field value > Use RAG Status measures. Map Green=#27AE60, Amber=#E67E22, Red=#C0392B."),
    ("Step 14: Test and validate",
     "Cross-check all KPI card totals against Reconciliation_Results.csv. Verify filter interactions. Test drill-through."),
    ("Step 15: Publish (optional)",
     "File > Publish to Power BI Service. Schedule data refresh if CSV files are in a shared location."),
]
for step, desc in steps:
    set_heading(doc2, step, 2)
    add_para(doc2, desc)

set_heading(doc2, "3. Dashboard Page Specifications", 1)
pages = [
    ("Page 1","Executive Financial Overview","KPI cards (10), waterfall chart, monthly trend, regional ranking, management exceptions, last refresh"),
    ("Page 2","Budget Performance","Approved vs revised, monthly profile, cumulative lines, absorption by directorate/CC/vote, underutilised budgets"),
    ("Page 3","Budget Review Workflow","Request counts, amounts, aging table, status funnel, turnaround metrics, action register"),
    ("Page 4","Commitments and Cash Exposure","Open commitments, aging, by category, expected payment schedule, available balance after commits"),
    ("Page 5","Procurement Monitoring","Plan vs contract vs paid, by method/category, delayed procurements, risk table"),
    ("Page 6","Pending Bills","Total, by aging band, by reason, by directorate, bills >90 days, payment forecast"),
    ("Page 7","Revenue and AIA","Budget vs actual, by stream/region, monthly trend, banking delays, unreconciled"),
    ("Page 8","Performance and Value for Money","KPI vs budget, cost per output, 4-quadrant scatter plot, achievement table"),
    ("Page 9","Forecasting and Financial Risk","YE forecast, scenarios, required monthly spend, risk heatmap, forecast methodology"),
    ("Page 10","Audit, Controls and Data Quality","Exceptions by type, amount at risk, repeat findings, DQ score, reconciliation status"),
    ("Page 11","Cost Centre Drill-Through","Full CC detail: budget, exp, commits, requests, bills, procurement, PI, exceptions, forecast"),
    ("Page 12","Management Action Register","Searchable table of all actions with status, responsible role, target date, comments"),
]
add_table(doc2, ["Page","Title","Key Visuals"], pages)

set_heading(doc2, "4. RAG Conditional Formatting Rules", 1)
rag_rules = [
    ("Budget Absorption","Green: within ±10% of time-elapsed","Amber: 10-20% gap","Red: >20% gap"),
    ("Forecast Over-Expenditure","Green: ≤100% of budget","Amber: 100-105%","Red: >105%"),
    ("Request Aging","Green: 0-14 days","Amber: 15-30 days","Red: >30 days"),
    ("Pending Bills Aging","Green: 0-30 days","Amber: 31-90 days","Red: >90 days"),
    ("Data Quality Score","Green: ≥98%","Amber: 95-98%","Red: <95%"),
    ("Revenue Achievement","Green: ≥90%","Amber: 70-90%","Red: <70%"),
]
add_table(doc2, ["Metric","Green","Amber","Red"], rag_rules)

filepath2 = f"{BASE}/03_PowerBI_Build_Pack/PowerBI_Implementation_Guide.docx"
doc2.save(filepath2)
print(f"  Saved: {filepath2}")
print("Documents built successfully.")
