"""
Parastatal X - Management Analysis Documents Builder
Creates: Senior Finance Manager commentary, EDA Report,
         Management Action Plan (XLSX), Financial Risk Register (XLSX),
         Reconciliation Report (XLSX)
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import xlsxwriter
import pandas as pd
import os

from pathlib import Path
BASE = Path(__file__).resolve().parents[1]
DATA = f"{BASE}/01_Data"
OUT  = f"{BASE}/04_Analysis"

FOREST_GREEN = RGBColor(0x1B, 0x43, 0x32)
CHARCOAL     = RGBColor(0x2C, 0x3E, 0x50)
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)

def doc_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = FOREST_GREEN
        run.font.size = Pt(16 if level==1 else 13 if level==2 else 11)

def doc_para(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = CHARCOAL
    return p

# ═══════════════════════════════════════════════════════════════════════════
# SENIOR FINANCE MANAGER COMMENTARY
# ═══════════════════════════════════════════════════════════════════════════
print("Building Senior_Finance_Manager_Analysis.docx...")
doc = Document()
p = doc.add_heading("PARASTATAL X", 0)
p.runs[0].font.color.rgb = FOREST_GREEN
p.runs[0].font.size = Pt(20)
doc.add_heading("SENIOR FINANCE MANAGER – MANAGEMENT COMMENTARY", 1).runs[0].font.color.rgb = FOREST_GREEN
doc_para(doc, "Financial Year 2025/2026 | As at 31 May 2026 | All figures in Kenya Shillings", italic=True)
doc_para(doc, "DISCLAIMER: All data used in this commentary is entirely synthetic.", italic=True)
doc_para(doc, "")

doc_heading(doc, "1. EXECUTIVE SUMMARY", 1)
doc_para(doc, """As at 31 May 2026, Parastatal X had utilised KSh 1,478,365,800 — representing 
87.0% of the approved annual budget of KSh 1,700,000,000. Against the expected time-based 
utilisation rate of 91.7% (11 of 12 months elapsed), the Organisation is 4.7 percentage 
points below the expected absorption level. While the overall position is within the amber 
management threshold, significant performance disparities exist across directorates and cost 
centres, requiring immediate management attention.

The blended year-end forecast of KSh 1,763,490,330 — representing 103.7% of the approved 
budget — indicates a slight over-expenditure risk of KSh 63,490,330. This is primarily driven 
by a high volume of open commitments (KSh 700,000,000) and expenditure acceleration expected 
in June 2026, the final month of the financial year.

Management is required to take immediate action to rebalance expenditure, clear pending bills, 
accelerate underperforming cost centres, and put in place controls to prevent over-expenditure 
in high-risk vote lines.""")

doc_heading(doc, "2. BUDGET UTILISATION", 1)
doc_para(doc, """The approved annual budget of KSh 1,700,000,000 has been distributed across 32 cost 
centres spanning 13 directorates and 8 regional offices. As at the cut-off date of 31 May 2026:

• YTD Actual Expenditure: KSh 1,478,365,800 (87.0% of approved budget)
• Expected Utilisation Rate (time-based): 91.7%
• Absorption Gap: -4.7 percentage points

Of the 32 active cost centres, 23 are rated Red (below 80% absorption versus time elapsed), 
7 are Amber (80-90%), and only 2 are Green. Five cost centres account for over 42% of the 
projected underutilisation, all within the Projects and Infrastructure Directorate and 
Regional Operations Directorate. These areas have experienced procurement delays and 
implementation challenges that have constrained expenditure absorption.

Conversely, the Finance and Accounts Directorate, ICT Directorate, and Human Resources 
Directorate have maintained absorption rates above 90%, demonstrating strong budget 
management and timely processing of transactions.""")

doc_heading(doc, "3. EXPENDITURE PERFORMANCE AND TRENDS", 1)
doc_para(doc, """Monthly expenditure has followed the expected seasonal pattern, with Q3 
(January–March 2026) recording higher spend in line with mid-year procurement completions. 
April 2026 showed a temporary dip as several large infrastructure contracts awaited 
commissioning milestones before final payments could be processed.

Year-on-year comparison shows overall expenditure growth of approximately 10.3% compared 
to FY2024/2025 (KSh 1,338,255,800 actual), consistent with the 9.7% budget increase 
from KSh 1,550,000,000 to KSh 1,700,000,000.

The average monthly expenditure of KSh 134,396,891 implies a required June 2026 expenditure 
of approximately KSh 221,634,200 to meet the approved budget — approximately 65% above the 
monthly average. Management must ensure that the necessary procurement, payment processing 
and commitment liquidation activities are prioritised and completed before 30 June 2026.""")

doc_heading(doc, "4. COMMITMENT EXPOSURE", 1)
doc_para(doc, """Open commitments as at 31 May 2026 total KSh 700,000,000, representing 47.3% 
of the remaining approved budget after YTD expenditure. A significant proportion of these 
commitments are associated with the Projects and Infrastructure Directorate and the 
Procurement and Supply Chain Directorate.

Key commitment concerns include:
• 18 cost centres with over-expenditure risk when commitments are netted against available balance
• Open commitments outstanding for more than 90 days requiring urgent management follow-up
• Several construction and equipment contracts where delivery timelines have slipped, creating 
  uncertainty over year-end payment timing

Management should review all open commitments above KSh 5,000,000 individually and determine 
which can be paid before year-end and which require carry-forward treatment.""")

doc_heading(doc, "5. PENDING BILLS", 1)
doc_para(doc, """The pending bills register as at 31 May 2026 records KSh 1,121,712,000 in 
outstanding invoices across all financial years. Of this:

• FY2025/2026 pending bills: KSh 1,121,712,000
• Bills aged over 90 days: The most significant concern requiring immediate clearance
• Primary reasons: Incomplete documentation (28%), pending approval (22%), insufficient 
  budget (18%), cash-flow constraints (15%)

The Organisation must prioritise settlement of invoices outstanding for more than 90 days 
to avoid escalating audit concerns and supplier relationship deterioration. A payment schedule 
should be prepared and approved by the Director, Finance and Accounts, with weekly progress 
reporting to the Chief Executive Officer.""")

doc_heading(doc, "6. BUDGET REVIEW AND APPROVAL WORKFLOW", 1)
doc_para(doc, """The budget review and approval process has experienced bottlenecks during the 
financial year. As at 31 May 2026, 1,000 requests have been received across all financial 
years, with a significant proportion remaining in the processing pipeline:

• Total approved amount: KSh 3,248,154,000 (FY2025/2026)
• Total processed amount: KSh 770,368,000
• Outstanding approved but unprocessed: KSh 2,477,786,000
• Requests outstanding more than 30 days: A significant proportion requiring escalation

The average approval turnaround time suggests processing bottlenecks at the Head of Budget 
level. Management should strengthen processing capacity, implement a five-day processing 
target, and introduce weekly reporting on unprocessed requests to the Director, Finance and 
Accounts.""")

doc_heading(doc, "7. REVENUE AND APPROPRIATIONS-IN-AID", 1)
doc_para(doc, """Revenue collection as at 31 May 2026 stands at KSh 979,529,000 against a 
budgeted target of KSh 1,110,412,000, representing an achievement rate of 88.2% (KSh 
130,883,000 below target).

Key revenue observations:
• Service Fees and Licensing Fees are the highest-performing revenue streams
• Training Fees and Consultancy Income are underperforming against targets
• Coast Region and North Eastern Region are experiencing the largest revenue shortfalls
• Banking delays in excess of 5 days have been flagged on several regional collections

Revenue managers must intensify collection efforts in June 2026, issue outstanding demand 
notices, and ensure all collections are banked within the prescribed 5-day window.""")

doc_heading(doc, "8. FORECAST YEAR-END POSITION", 1)
doc_para(doc, """Based on the blended four-method forecast model:

METHOD                          | YEAR-END FORECAST | ABSORPTION
Straight-Line Run Rate          | KSh 1,612,762,691 | 94.9%
3-Month Moving Average          | KSh 1,653,229,667 | 97.2%
Seasonal Adjustment             | KSh 1,722,723,783 | 101.3%
Commitment-Adjusted             | KSh 2,178,365,800 | 128.1%
BLENDED SELECTED FORECAST       | KSh 1,763,490,330 | 103.7%

The blended forecast of KSh 1,763,490,330 (103.7%) presents an over-expenditure risk of 
KSh 63,490,330 against the approved budget. Under the optimistic scenario (KSh 1,710,585,620), 
the Organisation would still marginally exceed the approved ceiling by KSh 10,585,620, 
suggesting that even with improved implementation, over-expenditure risk cannot be fully 
eliminated without immediate management controls.

Three scenarios are presented:
• Base Case: KSh 1,763,490,330 — over-expenditure of KSh 63.5M
• Optimistic: KSh 1,710,585,620 — over-expenditure of KSh 10.6M  
• Adverse: KSh 1,904,569,556 — over-expenditure of KSh 204.6M""")

doc_heading(doc, "9. MAJOR RISKS", 1)
risks = [
    ("RISK 1 – FORECAST OVER-EXPENDITURE","HIGH",
     "The blended forecast exceeds approved budget by KSh 63.5M. Five vote lines are "
     "approaching exhaustion. Immediate expenditure controls and reallocation from "
     "underspent votes are required."),
    ("RISK 2 – HIGH COMMITMENT EXPOSURE","HIGH",
     "KSh 700M in open commitments creates cash exposure that may crystallise before "
     "year-end. Several commitments are overdue and may require emergency payment."),
    ("RISK 3 – REVENUE UNDERPERFORMANCE","MEDIUM",
     "Revenue is KSh 130.9M (11.8%) below target. Continued underperformance will "
     "reduce AIA funding available in Q4 and may trigger cash-flow pressures."),
    ("RISK 4 – PROCUREMENT DELAYS","MEDIUM",
     "Multiple contracts are delayed by 30-120 days. If not completed by 30 June 2026, "
     "expenditure will be understated and activities will not be delivered."),
    ("RISK 5 – DATA AND AUDIT READINESS","MEDIUM",
     "Outstanding audit exceptions and missing supporting documents reduce audit "
     "readiness. All exceptions should be resolved before the financial year closes."),
]
for r_title, r_level, r_desc in risks:
    doc_para(doc, f"{r_title} (Risk Level: {r_level})", bold=True)
    doc_para(doc, r_desc)
    doc_para(doc, "")

doc_heading(doc, "10. RECOMMENDED MANAGEMENT ACTIONS", 1)
actions = [
    "Instruct the Director, Finance and Accounts to convene an emergency year-end expenditure "
    "review meeting by 10 June 2026, covering all cost centres rated Red or Amber.",
    "Issue a CEO directive restricting further expenditure in vote lines within 5% of "
    "their approved ceilings, effective immediately.",
    "Direct the Head of Budget to clear all unprocessed AIE requests within 5 working days.",
    "Require Regional Managers to submit and implement revised Q4 work plans by 12 June 2026.",
    "Instruct the Chief Accountant to prepare a payment schedule for all pending bills "
    ">90 days and seek Board approval for clearance.",
    "Direct the Revenue Manager to implement an emergency revenue collection drive for June 2026.",
    "Require the Head of Internal Audit to report on the status of all critical and high-risk "
    "audit exceptions by 15 June 2026.",
    "Prepare a supplementary budget request for Board consideration if forecast over-expenditure "
    "cannot be reduced below KSh 30M through internal management actions.",
]
for i, action in enumerate(actions, 1):
    doc_para(doc, f"{i}. {action}")

doc_para(doc, "")
doc_para(doc, "Prepared by: Senior Finance Manager, Parastatal X", italic=True)
doc_para(doc, "Date: June 2026 | Classification: Management Restricted | Synthetic Data Only", italic=True)

doc.save(f"{OUT}/Senior_Finance_Manager_Analysis.docx")
print("  Senior_Finance_Manager_Analysis.docx saved.")

# ═══════════════════════════════════════════════════════════════════════════
# EDA REPORT
# ═══════════════════════════════════════════════════════════════════════════
print("Building Exploratory_Data_Analysis_Report.docx...")
doc3 = Document()
doc3.add_heading("Parastatal X – Exploratory Data Analysis Report", 0).runs[0].font.color.rgb = FOREST_GREEN
doc_para(doc3, "FY2023/2024 to FY2025/2026 | Cut-off: 31 May 2026 | Synthetic Data", italic=True)
doc_para(doc3, "")

doc_heading(doc3, "1. Dataset Overview", 1)
eda_stats = [
    ("Approved Annual Budget",       "96 rows",    "3 FYs × 32 cost centres",         "KSh 4,650,000,000 total across 3 FYs"),
    ("Monthly Budget Allocation",    "1,152 rows", "3 FYs × 32 CCs × 12 months",      "Reconciles to annual approved budget"),
    ("Supplementary Budget",         "209 rows",   "60-90 per FY",                     "Mix of increases (65%) and reductions (35%)"),
    ("Budget Reallocations",         "224 rows",   "60-100 per FY",                    "75% approved; 25% pending/rejected"),
    ("Expenditure Transactions",     "12,100 rows","~4,033 per FY on average",         "KSh 4,134,987,100 total across 3 FYs"),
    ("Commitments",                  "1,350 rows", "~450 per FY",                      "Open commitments: KSh 700M (FY25/26)"),
    ("Budget Review Requests",       "1,000 rows", "~333 per FY",                      "1,000 requests; significant backlog identified"),
    ("Procurement Plan",             "390 rows",   "~130 per FY",                      "12 contracts delayed >45 days"),
    ("Pending Bills",                "520 rows",   "~173 per FY",                      "KSh 1,121,712,000 outstanding (FY25/26)"),
    ("Revenue and AIA",              "1,140 rows", "~380 per FY",                      "KSh 979,529,000 actual vs KSh 1,110,412,000 budget"),
    ("Performance Indicators",       "165 rows",   "15 indicators × 3 FYs × 3-4 qtrs","Average achievement: ~88%"),
    ("Audit Control Exceptions",     "300 rows",   "~100 per FY",                      "~35% resolved; 20% repeat findings"),
    ("Data Quality Log",             "250 rows",   "Across all datasets",              "~65% resolved; 18 Critical issues"),
]

def make_table_doc(doc, headers, rows):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True; run.font.size = Pt(9); run.font.color.rgb = WHITE
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1B4332"); shd.set(qn("w:val"), "clear")
        cell._tc.get_or_add_tcPr().append(shd)
    for r_idx, row in enumerate(rows):
        rc = table.rows[r_idx+1].cells
        bg = "F2F2F2" if r_idx%2==0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            rc[c_idx].text = str(val)
            for run in rc[c_idx].paragraphs[0].runs:
                run.font.size = Pt(9)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), bg); shd.set(qn("w:val"), "clear")
            rc[c_idx]._tc.get_or_add_tcPr().append(shd)

make_table_doc(doc3, ["Dataset","Volume","Coverage","Key Finding"], eda_stats)
doc_para(doc3, "")

doc_heading(doc3, "2. Key Positive Findings", 1)
positives = [
    ("Budget reconciliation accuracy","All monthly allocations reconcile exactly to approved annual budgets across all 3 financial years. Zero reconciliation difference. Demonstrates strong budget management discipline."),
    ("FY2025/2026 expenditure growth","YTD expenditure of KSh 1,478,365,800 represents 10.3% year-on-year growth from FY2024/2025 (KSh 1,338,255,800). This is consistent with the budget growth trajectory and organisational expansion."),
    ("Revenue stream diversification","Eight revenue streams are active across all 8 regions. Service Fees and Licensing Fees together account for ~58% of actual revenue collections, providing a diversified revenue base."),
    ("Commitment payment discipline","The majority of commitments (>75%) are within normal payment timelines. The commitment register is actively maintained with clear status tracking."),
    ("High-performing directorates","Finance and Accounts, ICT, and HR Directorates consistently achieve >90% absorption, demonstrating strong budget execution in headquarters functions."),
]
for title, detail in positives:
    doc_para(doc3, f"✓ {title}", bold=True)
    doc_para(doc3, detail)
    doc_para(doc3, "")

doc_heading(doc3, "3. Key Financial Risks Identified", 1)
risk_data = [
    ("Over-expenditure forecast","KSh 63,490,330","High","Projects & Infrastructure absorbing too fast in some lines while others lag"),
    ("Open commitment exposure","KSh 700,000,000","High","Procurement-heavy directorates; year-end payment pressure"),
    ("Revenue underperformance","KSh -130,883,000","Medium","Training and consultancy income below target"),
    ("Pending bills backlog","KSh 1,121,712,000","High","Documentation and budget insufficiency are top reasons"),
    ("Unprocessed AIE requests","KSh 2,477,786,000","High","Processing bottleneck at budget office level"),
]
make_table_doc(doc3, ["Risk","Amount","Level","Root Cause"], risk_data)
doc_para(doc3, "")

doc_heading(doc3, "4. Absorption Analysis by Directorate", 1)
dir_abs = [
    ("DIR-02 Finance and Accounts","92.1%","Green","On track"),
    ("DIR-04 Human Resource","91.5%","Green","On track"),
    ("DIR-06 ICT","93.8%","Green","Slight over-absorption risk"),
    ("DIR-03 Strategy and Planning","85.2%","Amber","Monitor closely"),
    ("DIR-09 Communications","83.7%","Amber","Delayed stakeholder events"),
    ("DIR-12 Projects and Infrastructure","74.1%","Red","Procurement delays"),
    ("DIR-13 Regional Operations","71.3%","Red","Activity implementation delays"),
    ("DIR-10 Research and M&E","68.9%","Red","Research contracts not yet awarded"),
    ("DIR-11 Technical Operations","70.2%","Red","Equipment procurement delays"),
    ("DIR-05 Procurement","79.8%","Amber","Good but under expected"),
]
make_table_doc(doc3, ["Directorate","Absorption %","RAG","Comment"], dir_abs)
doc_para(doc3, "")

doc_heading(doc3, "5. Data Quality Summary", 1)
dq_summary = [
    ("Missing Cost Centre Code","19","High","All resolved — mapped to default CC"),
    ("Duplicate Transaction Reference","14","Critical","Deduplicated; 14 duplicates removed"),
    ("Incorrect Vote Code","23","High","Corrected to valid vote codes"),
    ("Processed Amount > Approved","8","Critical","Capped at approved amount"),
    ("Invalid Date","11","High","Corrected transposition errors"),
    ("Inconsistent Region Names","31","Medium","Standardised to region codes"),
    ("Blank Approval Fields","27","Medium","Populated from workflow records"),
    ("Negative Expenditure Amounts","6","Critical","Sign corrected"),
    ("Leading/Trailing Spaces","45","Low","Trimmed in Power Query"),
    ("Different Capitalisation","38","Low","Standardised to uppercase codes"),
]
make_table_doc(doc3, ["Issue Type","Count","Severity","Resolution"], dq_summary)
doc_para(doc3, f"\nData quality score: 96.2% (250 issues in 9,500+ records = 2.6% error rate). "
         "This is within the acceptable 2-5% threshold for raw data demonstration. "
         "All issues are logged in Data_Quality_Log.csv and a cleaned version is available.")

doc3.save(f"{OUT}/Exploratory_Data_Analysis_Report.docx")
print("  Exploratory_Data_Analysis_Report.docx saved.")

# ═══════════════════════════════════════════════════════════════════════════
# MANAGEMENT ACTION PLAN (XLSX)
# ═══════════════════════════════════════════════════════════════════════════
print("Building Management_Action_Plan.xlsx...")
wb = xlsxwriter.Workbook(f"{OUT}/Management_Action_Plan.xlsx")
ws = wb.add_worksheet("Action_Plan")

FGRN = "#1B4332"; GOLD = "#B7950B"; LGRY = "#F2F2F2"; CHRC = "#2C3E50"
RED_B = "#FADBD8"; AMB_B = "#FFF2CC"; GRN_B = "#D5E8D4"

hdr_f = wb.add_format({'bold':True,'bg_color':FGRN,'font_color':'#FFFFFF','border':1,'font_size':10,'align':'center','valign':'vcenter','text_wrap':True})
body_f = wb.add_format({'font_color':CHRC,'border':1,'font_size':9,'valign':'vcenter','text_wrap':True})
grey_f = wb.add_format({'font_color':CHRC,'border':1,'font_size':9,'bg_color':LGRY,'valign':'vcenter','text_wrap':True})
crit_f = wb.add_format({'font_color':CHRC,'border':1,'font_size':9,'bg_color':RED_B,'valign':'vcenter','text_wrap':True})
high_f = wb.add_format({'font_color':CHRC,'border':1,'font_size':9,'bg_color':AMB_B,'valign':'vcenter','text_wrap':True})
good_f = wb.add_format({'font_color':CHRC,'border':1,'font_size':9,'bg_color':GRN_B,'valign':'vcenter','text_wrap':True})
pct_f  = wb.add_format({'font_color':CHRC,'border':1,'num_format':'0%','font_size':9,'valign':'vcenter'})
date_f = wb.add_format({'font_color':CHRC,'border':1,'num_format':'DD-MMM-YYYY','font_size':9,'valign':'vcenter'})
cur_f  = wb.add_format({'font_color':CHRC,'border':1,'num_format':'#,##0','font_size':9,'valign':'vcenter'})

headers = ["ActionID","Finding","RiskLevel","RecommendedAction","ResponsibleRole",
           "Priority","StartDate","TargetDate","ExpectedFinancialImpact",
           "ExpectedOperationalImpact","Status","ProgressPct","FollowUpDate","Comments"]
widths  = [10,45,12,50,35,10,14,14,25,35,14,12,14,40]
for c, (h,w) in enumerate(zip(headers,widths)):
    ws.write(0,c,h,hdr_f); ws.set_column(c,c,w)
ws.set_row(0,35)

actions = [
    ("ACT-001","Forecast over-expenditure of KSh 63.5M across 5 vote lines","High",
     "Impose immediate expenditure controls on at-risk vote lines; reallocate from underspent votes; seek supplementary budget if needed",
     "Director, Finance and Accounts","High","2026-06-05","2026-06-25","Prevent KSh 63.5M over-expenditure",
     "Maintain fiscal discipline; avoid audit queries","In Progress",0.40,"2026-06-15",
     "Vote line controls implemented for 3 lines; reallocation request submitted"),
    ("ACT-002","KSh 700M open commitments — over-expenditure risk","High",
     "Review all open commitments individually; clear overdue ones before year-end; cancel where goods/services not received",
     "Director, Supply Chain Management","High","2026-06-05","2026-06-25","Reduce commitment exposure by KSh 200M",
     "Clean procurement register for new FY","In Progress",0.30,"2026-06-12",
     "8 commitments escalated to Director level; 3 cancelled"),
    ("ACT-003","Revenue 11.8% below target (KSh -130.9M)","Medium",
     "Implement June 2026 revenue drive; issue demand notices; deploy regional revenue teams",
     "Revenue Manager","Medium","2026-06-05","2026-06-30","Recover KSh 60-80M in June",
     "Improve revenue achievement to 94%+","Pending",0.00,"2026-06-15",
     "Revenue drive plan submitted to CEO for approval"),
    ("ACT-004","23 cost centres with Red RAG absorption status","Critical",
     "Weekly reporting to CEO; revised work plans; accelerate pending procurements",
     "Chief Executive Officer","Critical","2026-06-03","2026-06-30","Improve absorption from 87% to 95%+",
     "Deliver planned activities and outputs","In Progress",0.15,"2026-06-10",
     "CEO directive issued; bi-weekly reporting started"),
    ("ACT-005","Pending bills >90 days — audit risk","High",
     "Prepare payment schedule; prioritise clearance of all bills >90 days before year-end",
     "Chief Accountant","High","2026-06-05","2026-06-30","Clear bills outstanding >90 days",
     "Reduce audit queries on pending payments","In Progress",0.50,"2026-06-10",
     "Payment schedule approved; 23 invoices queued for payment"),
    ("ACT-006","Budget review request backlog (KSh 2.48B unprocessed)","High",
     "Deploy additional processing staff; implement 5-day turnaround target; escalate stale requests",
     "Head of Budget","High","2026-06-05","2026-06-20","Clear KSh 500M backlog by 20 June",
     "Unblock cost centres from executing activities","In Progress",0.35,"2026-06-10",
     "2 additional finance officers deployed to budget unit"),
    ("ACT-007","12 procurement contracts delayed","Medium",
     "Issue contract extensions; re-procure where necessary; report to Board on impact",
     "Director, Supply Chain Management","Medium","2026-06-08","2026-06-30","Recover KSh 150M in payments by year-end",
     "Complete contracted activities before FY close","Pending",0.00,"2026-06-15",
     "Contract review meeting scheduled for 15 June 2026"),
    ("ACT-008","Audit exceptions — missing documents","Medium",
     "Circular to all cost centres; 72-hour deadline for document submission; follow up with Internal Audit",
     "Head of Internal Audit","Medium","2026-06-05","2026-06-20","Reduce audit exceptions by 60%",
     "Improve audit readiness score","Pending",0.10,"2026-06-12",
     "Circular issued; 35% of documents received to date"),
    ("ACT-009","Donor programme underspend","High",
     "Accelerate donor activities; engage donor on utilisation timeline; submit revised implementation plan",
     "Project Manager","High","2026-06-05","2026-06-30","Increase donor utilisation by KSh 80M",
     "Maintain donor confidence and programme continuity","Pending",0.00,"2026-06-15",
     "Donor notified; revised plan under preparation"),
    ("ACT-010","Data quality exceptions — 35 unresolved critical/high","Medium",
     "Data Officer to resolve all critical exceptions within 5 days; update DQ log",
     "Data Officer","Medium","2026-06-05","2026-06-15","Achieve 98%+ DQ score",
     "Improve data reliability for year-end reporting","In Progress",0.40,"2026-06-10",
     "14 of 35 exceptions resolved; on track"),
]

for r_idx, row in enumerate(actions, 1):
    priority = row[5]
    fmt = crit_f if priority=="Critical" else high_f if priority=="High" else (grey_f if r_idx%2==0 else body_f)
    for c_idx, val in enumerate(row):
        if c_idx in [6,7,12]:
            ws.write(r_idx, c_idx, val, wb.add_format({'font_color':CHRC,'border':1,'font_size':9,'valign':'vcenter',
                                                         'bg_color':RED_B if priority=="Critical" else AMB_B if priority=="High" else (LGRY if r_idx%2==0 else "#FFFFFF")}))
        elif c_idx == 11:
            ws.write_number(r_idx, c_idx, val, wb.add_format({'font_color':CHRC,'border':1,'num_format':'0%','font_size':9,'valign':'vcenter',
                                                                'bg_color':RED_B if priority=="Critical" else AMB_B if priority=="High" else (LGRY if r_idx%2==0 else "#FFFFFF")}))
        else:
            ws.write(r_idx, c_idx, val, fmt)
    ws.set_row(r_idx, 40)

ws.freeze_panes(1,0)
wb.close()
print("  Management_Action_Plan.xlsx saved.")

# ═══════════════════════════════════════════════════════════════════════════
# FINANCIAL RISK REGISTER (XLSX)
# ═══════════════════════════════════════════════════════════════════════════
print("Building Financial_Risk_Register.xlsx...")
wb2 = xlsxwriter.Workbook(f"{OUT}/Financial_Risk_Register.xlsx")
ws2 = wb2.add_worksheet("Risk_Register")

hf = wb2.add_format({'bold':True,'bg_color':FGRN,'font_color':'#FFFFFF','border':1,'font_size':10,'align':'center','valign':'vcenter','text_wrap':True})
bf = wb2.add_format({'font_color':CHRC,'border':1,'font_size':9,'valign':'vcenter','text_wrap':True})
gf = wb2.add_format({'font_color':CHRC,'border':1,'font_size':9,'bg_color':LGRY,'valign':'vcenter','text_wrap':True})
cr = wb2.add_format({'font_color':CHRC,'border':1,'font_size':9,'bg_color':RED_B,'valign':'vcenter','text_wrap':True})
hi = wb2.add_format({'font_color':CHRC,'border':1,'font_size':9,'bg_color':AMB_B,'valign':'vcenter','text_wrap':True})
me = wb2.add_format({'font_color':CHRC,'border':1,'font_size':9,'bg_color':'#FDFEFE','valign':'vcenter','text_wrap':True})

rr_hdrs = ["RiskID","Category","RiskDescription","Likelihood","Impact","RiskRating","AmountAtRisk",
           "RootCause","ControlMeasures","ResidualRisk","ResponsibleRole","ReviewDate","Status","Comments"]
rr_widths = [10,18,45,12,12,12,18,35,35,14,30,14,14,35]
for c,(h,w) in enumerate(zip(rr_hdrs,rr_widths)):
    ws2.write(0,c,h,hf); ws2.set_column(c,c,w)
ws2.set_row(0,35)

risks = [
    ("RSK-001","Budget Overrun","Forecast over-expenditure of KSh 63.5M against approved budget","High","High","Critical","63,490,330",
     "Higher than expected Q4 expenditure; year-end commitment liquidation pressure",
     "Vote line expenditure controls; reallocation from underspent votes; supplementary budget request",
     "High","Director, Finance and Accounts","2026-06-10","Open",
     "CEO directive issued; monitoring weekly"),
    ("RSK-002","Commitment Crystallisation","KSh 700M open commitments may crystallise before year-end creating over-expenditure","High","High","Critical","700,000,000",
     "Large procurement contracts; delayed delivery; year-end payment timing uncertainty",
     "Commitment review; cancellation of stale orders; phased payment scheduling",
     "High","Director, Supply Chain Management","2026-06-10","Open",
     "Individual review of all commitments >KSh 5M initiated"),
    ("RSK-003","Revenue Underperformance","Revenue KSh 130.9M below target; AIA funding shortfall","Medium","High","High","130,883,000",
     "Lower than expected service demand; delayed licencing renewals; collection inefficiencies",
     "Revenue drive; demand notices; regional collection monitoring",
     "Medium","Revenue Manager","2026-06-15","Open",
     "June revenue drive approved by management"),
    ("RSK-004","Pending Bills Accumulation","Uncleared invoices totalling KSh 1.1B creating supplier and audit risk","High","Medium","High","1,121,712,000",
     "Incomplete documentation; budget insufficiency; approval delays",
     "Priority payment schedule; documentation drive; vote line controls",
     "Medium","Chief Accountant","2026-06-15","In Progress",
     "Payment schedule approved; 23 invoices queued"),
    ("RSK-005","Procurement Delays","12 contracts delayed >45 days; risk of undelivered activities","Medium","Medium","High","150,000,000",
     "Delayed tender processes; slow supplier mobilisation; specification changes",
     "Contract extensions; re-procurement; performance bond activation",
     "Medium","Director, Supply Chain Management","2026-06-20","Open",
     "Contract review meeting 15 June 2026"),
    ("RSK-006","Audit Exceptions","300 audit exceptions; 20% repeat findings; KSh 45M amount at risk","Medium","High","High","45,000,000",
     "Missing documents; incorrect coding; segregation of duties gaps",
     "Document drive; training; supervisory reviews; corrective action plans",
     "Medium","Head of Internal Audit","2026-06-20","In Progress",
     "Circular issued; 35% documents received"),
    ("RSK-007","Low Performance Absorption","23 cost centres at <80% absorption vs time elapsed","High","Medium","High","221,634,200",
     "Activity implementation delays; procurement bottlenecks; staff capacity",
     "Revised work plans; accelerated procurement; management oversight",
     "Medium","Chief Executive Officer","2026-06-15","In Progress",
     "CEO directive; bi-weekly reporting started"),
    ("RSK-008","Donor Programme Risk","Donor fund underutilisation may trigger clawback or programme disruption","Low","High","Medium","80,000,000",
     "Implementation delays; procurement non-compliance; reporting gaps",
     "Revised implementation plan; donor engagement; compliance review",
     "Low","Project Manager","2026-06-20","Open",
     "Donor notified; revised plan under preparation"),
    ("RSK-009","Data Quality Risk","2.6% data error rate; 35 critical/high exceptions unresolved","Low","Medium","Medium","25,000,000",
     "Data entry errors; system integration gaps; inadequate validation",
     "Data quality log monitoring; Power Query validation; staff training",
     "Low","Data Officer","2026-06-15","In Progress",
     "14 of 35 exceptions resolved"),
    ("RSK-010","Unprocessed AIE Backlog","KSh 2.48B in unprocessed approved requests blocking expenditure","High","Medium","High","2,477,786,000",
     "Processing capacity constraints; documentation gaps; staff workload",
     "Additional processing staff; 5-day SLA; escalation protocol",
     "Medium","Head of Budget","2026-06-20","In Progress",
     "2 additional finance officers deployed"),
]
for r_idx, row in enumerate(risks,1):
    rating = row[5]
    fmt = cr if rating=="Critical" else hi if rating=="High" else (gf if r_idx%2==0 else bf)
    for c_idx, val in enumerate(row):
        ws2.write(r_idx, c_idx, val, fmt)
    ws2.set_row(r_idx, 40)
ws2.freeze_panes(1,0)
wb2.close()
print("  Financial_Risk_Register.xlsx saved.")

# ═══════════════════════════════════════════════════════════════════════════
# RECONCILIATION REPORT (XLSX)
# ═══════════════════════════════════════════════════════════════════════════
print("Building Reconciliation_Report.xlsx...")
rec_src = f"{OUT}/Reconciliation_Results.csv"
if os.path.exists(rec_src):
    df_rec = pd.read_csv(rec_src)
    wb3 = xlsxwriter.Workbook(f"{OUT}/Reconciliation_Report.xlsx")
    ws3 = wb3.add_worksheet("Reconciliation_Report")
    hf3 = wb3.add_format({'bold':True,'bg_color':FGRN,'font_color':'#FFFFFF','border':1,'font_size':10,'align':'center','valign':'vcenter'})
    bf3 = wb3.add_format({'font_color':CHRC,'border':1,'font_size':9,'valign':'vcenter'})
    gf3 = wb3.add_format({'font_color':CHRC,'border':1,'font_size':9,'bg_color':LGRY,'valign':'vcenter'})
    cur3= wb3.add_format({'font_color':CHRC,'border':1,'num_format':'#,##0','font_size':9,'valign':'vcenter'})
    cur3g=wb3.add_format({'font_color':CHRC,'border':1,'num_format':'#,##0','font_size':9,'bg_color':LGRY,'valign':'vcenter'})
    ok3 = wb3.add_format({'font_color':'#1B4332','border':1,'font_size':9,'bg_color':GRN_B,'valign':'vcenter','bold':True})
    var3= wb3.add_format({'font_color':CHRC,'border':1,'font_size':9,'bg_color':AMB_B,'valign':'vcenter'})

    cols = list(df_rec.columns)
    widths3 = [55,20,20,18,18,22,50]
    for c,(h,w) in enumerate(zip(cols,widths3[:len(cols)])):
        ws3.write(0,c,h,hf3); ws3.set_column(c,c,w)
    ws3.set_row(0,30)

    for r_idx, row in df_rec.iterrows():
        is_grey = r_idx % 2 == 0
        status = str(row.get("Status",""))
        for c_idx, col in enumerate(cols):
            val = row[col]
            is_cur = col in ["SourceTotal","ModelTotal","Difference","AbsDifference"]
            if status == "Reconciled":
                fmt = ok3
            elif is_grey:
                fmt = var3
            else:
                fmt = bf3
            if is_cur:
                try:
                    ws3.write_number(r_idx+1, c_idx, float(val), fmt)
                except:
                    ws3.write(r_idx+1, c_idx, str(val), fmt)
            else:
                ws3.write(r_idx+1, c_idx, str(val), fmt)
        ws3.set_row(r_idx+1, 22)
    ws3.freeze_panes(1,0)
    wb3.close()
    print("  Reconciliation_Report.xlsx saved.")
else:
    print("  Reconciliation_Results.csv not found — skipping Reconciliation_Report.xlsx")

print("\nAll analysis documents built.")
